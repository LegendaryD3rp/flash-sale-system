#!/usr/bin/env python3
"""Part 1 - Cover + Abstract + Ch1-4"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()
for s in doc.sections:
    s.page_width = Cm(21.0); s.page_height = Cm(29.7)
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

sty = doc.styles['Normal']
sty.font.name = 'SimSun'; sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 1.5

H = doc.add_heading

def P(t, b=False, s=12, a='left', c=None):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.size = Pt(s); r.font.name = 'SimSun'
    if b: r.bold = True
    if c: r.font.color.rgb = c
    p.alignment = {'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}.get(a, WD_ALIGN_PARAGRAPH.LEFT)
    return p

def T(hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hdrs):
        c=t.rows[0].cells[i]; c.text=h
        for p2 in c.paragraphs:
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p2.runs: r.bold=True
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)

def B(t):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.size=Pt(12); r.font.name='SimSun'; r.bold=True
    return p

def C(t):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.name='Consolas'; r.font.size=Pt(9)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)

# ============ COVER ============
for _ in range(4): doc.add_paragraph()
P('课程期末大作业', s=28, b=True, a='center')
doc.add_paragraph()
P('AI驱动的"高并发秒杀与智能电商"后台系统', s=18, b=True, a='center', c=RGBColor(0,51,102))
doc.add_paragraph()
P('—— 闪购秒杀系统 (Flash Sale System)', s=14, a='center', c=RGBColor(100,100,100))
for _ in range(4): doc.add_paragraph()
for l,v in [('组  长：','X X X'),('联系方式：','xxx@xxx.edu.cn'),('组  员：','X X X, X X X'),
    ('代码仓库：','https://github.com/LegendaryD3rp/flash-sale-system'),
    ('提交日期：',datetime.date.today().strftime('%Y年%m月%d日'))]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(l).font.size=Pt(14); p.add_run(l).font.name='SimSun'; p.runs[0].bold=True
    p.add_run(v).font.size=Pt(14); p.add_run(v).font.name='SimSun'
doc.add_page_break()

# ============ ABSTRACT ============
H('摘要', level=1)
P('本系统为AI驱动的高并发秒杀与智能电商后台系统（题目一），实现了一个包含日常商品浏览、高并发秒杀活动、WebSocket实时推送和AI智能导购的微服务电商平台。系统采用Java 17 + Spring Boot 3.2.5 + Spring Cloud Gateway 2023.0.3作为技术底座，拆分为6个独立微服务和1个API网关，服务间通过HTTP RESTful接口通信。')
P('核心中间件：MySQL 8.0、Redis 7.0（库存预热与原子扣减）、RabbitMQ 3.12（异步削峰）、Spring Cloud Gateway（统一入口、JWT鉴权、限流）。系统以微服务架构实现高内聚低耦合。')
P('核心技术点：Gateway统一入口与自定义限流；Redis原子DECR实现10000并发0超卖；RabbitMQ异步订单削峰；WebSocket实时推送；LLM API集成AI导购。')
P('成果亮点：单商品10000请求7.4秒0错误；万人同场50品全覆盖10000单零错误；全流程E2E（注册->登录->地址->购物车->下单->支付->秒杀）10000人8步0错误。每个功能点均有自动化测试脚本覆盖，代码已推送至GitHub。')
doc.add_page_break()

# ============ CH1 ============
H('一、业务背景与范围界定', level=1)
H('1.1 业务场景', level=2)
P('秒杀是电商营销的核心手段，但瞬时流量可达日常百倍，对系统并发控制要求极高。本系统以高并发秒杀为核心场景，同时覆盖日常电商全流程，并引入AI大模型提供智能导购能力。')
H('1.2 用户角色与用例', level=2)
T(['角色','权限','核心用例'],[['普通用户','CUSTOMER','注册/登录、浏览商品、管理地址、购物车、下单/支付、秒杀抢购、AI咨询'],['管理员','ADMIN','管理商品、创建秒杀活动、查看所有订单、系统统计'],['AI导购','SERVICE','接收咨询->检索商品->调用LLM API->返回建议']])
H('1.3 MVP范围', level=2)
B('基础支撑功能：'); [doc.add_paragraph(i, style='List Bullet') for i in ['用户注册/登录，JWT签发与校验，角色鉴权','商品管理：增删改查、分页、库存','地址管理：CRUD','购物车管理：加购、列表、删除','订单管理：创建、支付、取消、列表']]
B('核心架构功能：'); [doc.add_paragraph(i, style='List Bullet') for i in ['统一API网关路由与JWT鉴权','Redis预热+原子扣库存防超卖','RabbitMQ异步订单削峰','WebSocket实时推送','Gateway限流（令牌桶）','AI智能导购（LLM API）']]
H('1.4 非目标范围', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in ['支付对接真实第三方（模拟支付）','分布式事务Seata（本地事务+MQ最终一致）','注册中心Nacos（直连IP）','容器化部署Docker/K8s','CDN/静态资源分离']]
doc.add_page_break()

# ============ CH2 ============
H('二、本地原型系统架构与数据流', level=1)
H('2.1 微服务清单', level=2)
T(['服务','端口','职责','依赖'],[['gateway','8080','API网关','无'],['user-service','8081','用户/登录/JWT','MySQL'],['product-service','8082','商品CRUD','MySQL'],['order-service','8083','订单/购物车/地址/WebSocket','MySQL,RMQ,Redis'],['ai-service','8084','AI导购LLM','外调LLM'],['seckill-service','8085','秒杀活动/预热/扣减','MySQL,Redis,RMQ']])
H('2.2 中间件清单', level=2)
T(['中间件','版本','用途'],[['MySQL','8.0','核心业务持久化'],['Redis','7.0','秒杀库存预热与原子扣减'],['RabbitMQ','3.12','异步削峰、订单消息'],['Spring Cloud Gateway','2023.0.3','统一入口、路由、鉴权、限流']])
H('2.3 部署拓扑', level=2)
C('Client -> API Gateway(:8080) [JWT Filter + RateLimit]\n  |--- user-service(:8081) -> MySQL\n  |--- product-service(:8082) -> MySQL\n  |--- order-service(:8083) -> MySQL + RabbitMQ + Redis\n  |--- seckill-service(:8085) -> MySQL + Redis + RabbitMQ\n  |--- ai-service(:8084) -> LLM API')
H('2.4 核心流程时序', level=2)
B('秒杀链：'); C('Client -> Gateway -> Seckill -> Redis(原子DECR) -> MQ -> Order(落库) -> WebSocket(推送给用户)')
B('普通购物链：'); C('Client -> Gateway -> User(注册登录) -> Order(地址/购物车/下单/支付)')
P('网关路由映射表：')
T(['路由前缀','目标服务','端口','说明'],[['/api/user/**,/api/address/**,/api/admin/user/**,/api/admin/audit-log/**','user-service','8081','用户、地址、管理后台用户管理'],['/api/product/**,/api/upload/**','product-service','8082','商品管理、文件上传'],['/api/seckill/**','seckill-service','8085','秒杀活动管理、秒杀抢购'],['/api/order/**,/api/admin/order/**,/api/cart/**,/api/coupon/**,/api/admin/coupon/**,/api/admin/statistics/**','order-service','8083','订单、购物车、优惠券、数据统计'],['/ws/**','order-service','8083（转发至flash-sale-common的SeckillWebSocketHandler）','WebSocket实时推送'],['/api/ai/**','ai-service','8084','AI智能导购']])
H('2.5 安全与网关策略', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in ['JWT白名单：公开接口不校验','Gateway JWT Filter验证非白名单请求','X-User-Id Header透传','RateLimitFilter：seckill-capacity=2000, general-capacity=4000']]
doc.add_page_break()

# ============ CH3 ============
H('三、数据库设计与持久化策略', level=1)
H('3.1 ER图', level=2)
C('user(id,username,password,role,status) --< address(user_id)\n  |--< cart(user_id)\n  |--< order(user_id)\nproduct(id,name,price,stock) --< cart(product_id)\n  |--< order(product_id)\n  |--< seckill_activity(product_id)\nseckill_activity(id,product_id,total_stock,status) --< order(seckill_activity_id)')
H('3.2 索引策略', level=2)
T(['表','索引','列','原因'],[['order','idx_user_id','user_id','我的订单查询'],['order','idx_status','status','按状态筛选'],['order','idx_seckill_activity','seckill_activity_id','活动关联订单'],['cart','idx_user_id','user_id','购物车查询'],['user','uk_username','username','唯一登录名']])
H('3.3 事务与一致性', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in ['普通下单：@Transactional + 死锁重试3次(50/100/200ms)','秒杀：Redis原子扣减(强一致)+MQ(最终一致)+幂等消费','支付状态校验：仅PENDING_PAY可支付']]
doc.add_page_break()

# ============ CH4 ============
H('四、缓存、消息与并发控制策略', level=1)
H('4.1 Redis使用', level=2)
T(['用途','数据结构','Key','说明'],[['秒杀库存','String','seckill:stock:{id}','预热SET，秒杀DECR'],['预热状态','String','seckill:warm:{id}','避免重复预热']])
C('预热：redis.set("seckill:stock:{id}", 200)\n秒杀：Long r = redis.decrement(key); if(r>=0)成功 else 售罄')
H('4.2 RabbitMQ', level=2)
T(['配置','值'],[['队列','seckill.order.queue'],['交换机','seckill.order.exchange(直连)'],['Prefetch','10'],['并发消费者','3-8'],['死信队列','重试3次失败后入DLQ']])
H('4.3 并发保护', level=2)
B('Gateway限流：'); T(['维度','参数'],[['通用','capacity=4000, refill=2000'],['秒杀','seckill-capacity=2000, refill=1000']])
B('服务层保护：'); [doc.add_paragraph(i, style='List Bullet') for i in ['死锁重试指数退避','Redis+MQ无DB写压力','连接池控制']]

doc.save('/tmp/report_part1.docx')
print('OK: /tmp/report_part1.docx')
