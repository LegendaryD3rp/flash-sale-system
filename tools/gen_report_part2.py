#!/usr/bin/env python3
"""Part 2 - Ch5-11 + merge"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Load part 1
doc = Document('/tmp/report_part1.docx')

def P(t, b=False, s=12, a='left', c=None):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.size=Pt(s); r.font.name='SimSun'
    if b: r.bold=True
    if c: r.font.color.rgb=c
    p.alignment={'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}.get(a,WD_ALIGN_PARAGRAPH.LEFT)
    return p

def B(t):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.size=Pt(12); r.font.name='SimSun'; r.bold=True; return p

def T(hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hdrs):
        c=t.rows[0].cells[i]; c.text=h
        for p2 in c.paragraphs:
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p2.runs: r.bold=True
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            t.rows[ri+1].cells[ci].text=str(val)

def C(t):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.name='Consolas'; r.font.size=Pt(9)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)

H = doc.add_heading

# ============ CH5 ============
doc.add_page_break()
H('五、AI/LLM 集成设计', level=1)
H('5.1 模型选型', level=2)
T(['配置','值'],[['模型','deepseek-chat / qwen-plus（可配置）'],['协议','OpenAI兼容 /v1/chat/completions'],['上下文','128K tokens'],['限流','约60 RPM']])
H('5.2 调用方式', level=2)
P('用户通过POST /api/ai/chat发送咨询，ai-service先从product-service获取商品列表拼接成system prompt，再调用LLM API获取回复。当前为同步调用（请求-响应模式）。')
C('请求: POST /api/ai/chat {"query":"推荐办公笔记本"}\n-> 获取商品列表拼入system prompt\n-> POST /v1/chat/completions\n-> 返回自然语言推荐')
H('5.3 生产优化方向', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in [
    '请求排队：Redis队列控制并发不超过API限流（60 RPM）',
    '相似缓存：缓存重复问题回复（语义相似度>0.95）',
    '流式推送：改为SSE逐步返回',
    '降级：API不可用时返回预设FAQ']]

# ============ CH6 ============
doc.add_page_break()
H('六、实时通信设计（WebSocket）', level=1)
H('6.1 架构', level=2)
P('客户端连接 ws://localhost:8080/ws/seckill（经Gateway路由到order-service）。订单创建成功后SeckillOrderConsumer通过SeckillWebSocketHandler推送消息到对应用户。')
C('Client -> ws://:8080/ws/seckill -> Gateway -> OrderService(:8083)\n  -> SeckillWebSocketHandler.sessionMap<userId, session>\n  -> 订单创建成功后推送: {"type":"ORDER_CREATED","orderId":...,"status":"SUCCESS"}')
H('6.2 跨JVM问题', level=2)
P('SeckillWebSocketHandler.sessionMap使用static ConcurrentHashMap，多实例部署时WebSocket连接仅落在单个节点，其他节点订单无法推送。')
P('解决方案：将Gateway的/ws/**路由从seckill-service改为order-service，同时添加白名单GET:/ws/。', b=True)
P('生产方案：集成Redis Pub/Sub或RabbitMQ广播，所有实例订阅同一个channel，收到推送事件后检查sessionMap是否有目标用户。', b=True)

# ============ CH7 ============
doc.add_page_break()
H('七、技术选型依据与权衡', level=1)
H('7.1 为什么选择Spring Cloud Gateway而非Nginx？', level=2)
P('Nginx配置变更需reload或重启，不适合频繁调整路由的微服务开发阶段。Gateway原生支持Java配置、动态路由、与JWT Filter无缝集成。生产环境可在Nginx前置做L4负载均衡。')
H('7.2 为什么选择Redis原生DECR而非Lua脚本？', level=2)
P('DECR是原子操作，单key递减天然线性化，无需Lua脚本的额外解析开销。Redis官方文档指出DECR时间复杂度O(1)，比Lua脚本的EVAL（需脚本缓存+KEYS解析）更适合超高并发的简单扣减场景。')
H('7.3 为什么选择RabbitMQ而非Kafka？', level=2)
P('Kafka吞吐量虽高（百万级/秒），但对消息确认（ACK）延迟敏感。秒杀场景每秒约1000-2000订单，RabbitMQ完全胜任，且提供DLQ+手动ACK精确保证"不丢不重"。')
H('7.4 为什么选择MyBatis-Plus而非JPA？', level=2)
P('JPA的自动DDL和懒加载在高并发场景下容易引发N+1查询。MyBatis-Plus提供分页插件、代码生成器，且Mapper XML可直接优化SQL执行计划。')
H('7.5 为什么使用WebSocket而非SSE？', level=2)
P('SSE为单向推送（服务器→客户端），秒杀场景需要双向通信（用户确认、重试等）。WebSocket是JSR 356标准，Spring原生支持，与HTTP长连接兼容。')

# ============ CH8 (NEW) ============
doc.add_page_break()
H('八、API接口设计', level=1)
P('系统共实现68个RESTful API端点，按微服务模块组织如下：', s=12)
H('8.1 用户服务（user-service）', level=2)
T(['方法','路径','说明','角色'],
[['POST','/api/user/register','用户注册','公开'],
 ['POST','/api/user/login','用户登录','公开'],
 ['GET','/api/user/me','获取当前用户信息','已登录'],
 ['PUT','/api/user/profile','修改个人资料','已登录'],
 ['PUT','/api/user/password','修改密码','已登录'],
 ['POST','/api/user/forgot-password','忘记密码（返回重置token）','公开'],
 ['POST','/api/user/reset-password','重置密码','公开'],
 ['GET','/api/address/list','查询地址列表','已登录'],
 ['POST','/api/address','新增地址','已登录'],
 ['PUT','/api/address/{id}','修改地址','已登录'],
 ['DELETE','/api/address/{id}','删除地址','已登录'],
 ['PUT','/api/address/{id}/default','设为默认地址','已登录'],
 ['GET','/api/admin/user/list','分页查询用户列表','ADMIN'],
 ['GET','/api/admin/user/{id}','查看用户详情','ADMIN'],
 ['PUT','/api/admin/user/{id}','修改用户信息','ADMIN'],
 ['DELETE','/api/admin/user/{id}','软删除/禁用用户','ADMIN'],
 ['GET','/api/admin/audit-log','查询审计日志','ADMIN']])
H('8.2 商品服务（product-service）', level=2)
T(['方法','路径','说明','角色'],
[['GET','/api/product','分页查询商品列表','公开'],
 ['GET','/api/product/search','搜索商品（名称+描述模糊匹配）','公开'],
 ['GET','/api/product/{id}','查询商品详情','公开'],
 ['POST','/api/product','新增商品','ADMIN'],
 ['PUT','/api/product/{id}','编辑商品','ADMIN'],
 ['DELETE','/api/product/{id}','删除商品','ADMIN'],
 ['POST','/api/product/review','新增评价','已登录'],
 ['POST','/api/product/{productId}/reviews','对指定商品发表评价','已登录'],
 ['GET','/api/product/{productId}/reviews','查询商品评价列表','公开'],
 ['GET','/api/product/review/check','检查某订单是否已评价','已登录'],
 ['GET','/api/product/review/mine','我的评价列表','已登录'],
 ['POST','/api/product/favorite','收藏商品','已登录'],
 ['DELETE','/api/product/favorite/{productId}','取消收藏','已登录'],
 ['GET','/api/product/favorite/favorites','我的收藏列表（分页）','已登录'],
 ['GET','/api/product/favorite/check','批量检查是否已收藏','已登录']])
H('8.3 订单服务（order-service）', level=2)
T(['方法','路径','说明','角色'],
[['GET','/api/order/list','当前用户订单列表','已登录'],
 ['GET','/api/order/{orderId}','订单详情（带归属校验）','已登录'],
 ['POST','/api/order/create','从购物车创建订单','已登录'],
 ['POST','/api/order/{id}/pay','支付订单（状态机推进）','已登录'],
 ['POST','/api/order/{id}/receive','确认收货','已登录'],
 ['POST','/api/order/{id}/cancel','取消订单（释放优惠券）','已登录'],
 ['GET','/api/admin/order/list','管理员查看所有订单','ADMIN'],
 ['POST','/api/admin/order/{id}/ship','管理员发货','ADMIN'],
 ['GET','/api/cart/list','查询用户购物车','已登录'],
 ['GET','/api/cart/listByIds','按ID列表查询购物车项','已登录'],
 ['POST','/api/cart/add','添加商品到购物车','已登录'],
 ['PUT','/api/cart/{id}','修改购物车项数量','已登录'],
 ['DELETE','/api/cart/{id}','删除一项','已登录'],
 ['POST','/api/cart/clear','清空购物车','已登录'],
 ['GET','/api/admin/coupon/list','优惠券列表','ADMIN'],
 ['GET','/api/admin/coupon/{id}','优惠券详情','ADMIN'],
 ['POST','/api/admin/coupon','新建优惠券','ADMIN'],
 ['PUT','/api/admin/coupon/{id}','编辑优惠券','ADMIN'],
 ['PUT','/api/admin/coupon/{id}/disable','下架优惠券','ADMIN'],
 ['POST','/api/coupon/claim/{couponId}','用户领取优惠券','已登录'],
 ['GET','/api/coupon/mine','查看我的优惠券','已登录'],
 ['GET','/api/coupon/available','查询可用优惠券','已登录'],
 ['POST','/api/coupon/calc-discount','计算优惠金额','已登录'],
 ['GET','/api/admin/statistics/summary','今日概况','ADMIN'],
 ['GET','/api/admin/statistics/order-trend','订单趋势','ADMIN'],
 ['GET','/api/admin/statistics/top-products','热销商品Top','ADMIN']])
H('8.4 秒杀服务（seckill-service）', level=2)
T(['方法','路径','说明','角色'],
[['GET','/api/seckill/activity/list','活动列表','公开'],
 ['GET','/api/seckill/activity/{id}','活动详情','公开'],
 ['POST','/api/seckill/activity','创建活动','ADMIN'],
 ['PUT','/api/seckill/activity/{id}','编辑活动','ADMIN'],
 ['PATCH','/api/seckill/activity/{id}/status','变更活动状态','ADMIN'],
 ['POST','/api/seckill/activity/{id}/warm-up','手动预热','ADMIN'],
 ['POST','/api/seckill/flash','秒杀抢购入口','已登录'],
 ['GET','/api/seckill/ping','心跳检测','公开']])
H('8.5 AI服务（ai-service）', level=2)
T(['方法','路径','说明','角色'],
[['POST','/api/ai/chat','AI导购对话','公开']])
H('8.6 公共模块（common）', level=2)
T(['方法','路径','说明','角色'],
[['POST','/api/upload','文件上传（仅图片）','ADMIN']])
H('8.7 秒杀接口请求/响应示例', level=2)
P('以秒杀抢购接口 POST /api/seckill/flash 为例：')
B('请求格式：')
C('POST /api/seckill/flash\nContent-Type: application/json\nX-User-Id: 1\n\n{\n  "activityId": 1\n}')
B('成功响应：')
C('{\n  "code": 200,\n  "message": "秒杀请求已受理",\n  "data": "排队中",\n  "timestamp": 1712345678000\n}')
B('失败响应（库存不足）：')
C('{\n  "code": 400,\n  "message": "库存不足",\n  "timestamp": 1712345678000\n}')
P('其中activityId为秒杀活动ID（@NotNull校验），userId由Gateway从JWT解析后通过X-User-Id请求头透传。后端通过Redis原子DECR扣减库存，再将订单创建消息投递至RabbitMQ异步落单，最终通过WebSocket推送抢购结果。')

# ============ CH9 (was CH8) ============
doc.add_page_break()
H('九、测试', level=1)
H('9.1 测试体系', level=2)
T(['层次','工具','范围','数量'],[['单元测试','JUnit 5 + Mockito','Java后端核心逻辑','88个全部通过'],['E2E测试','Python + aiohttp','HTTP全链路','53个(50通过3跳过)'],['架构验证','Python + aiohttp','架构9概念','11/11通过'],['万人演示','Python + aiohttp','10000并发','7个脚本'],['Postman','Postman Collection','核心业务链','一键导入按顺序执行']])
H('9.2 性能数据', level=2)
T(['场景','请求','耗时','成功','错误'],[['单商品1万秒杀','10,000','7.4s','10,000','0'],['50品万人同场(DB直插)','10,000','19.8s','9,694','0'],['50品万人同场(HTTP全流程)','10,000x8','565s','80,000步','0'],['全流程E2E 8步','10,000人','565s','10,000人','0']])
H('9.3 测试结果说明', level=2)
P('Java单元测试88个全部通过（用户服务36个、商品服务24个、订单服务21个、秒杀服务7个）。Python E2E测试共53个测试用例（41个单接口+8个E2E场景+5个并发安全），其中50个通过、3个因依赖环境跳过（Postman Collection也覆盖了全部核心业务链）。万人并发演示中，单商品秒杀7.4秒完成10000笔0错误，全流程E2E（注册→登录→选品→购物车→下单→支付）565秒完成10000人。')
H('9.4 并发正确性验证', level=2)
P('通过 stress_validation_v2.py 的11项架构验证测试，包括：库存不超卖验证、一人一单约束验证、订单状态机正确性验证、支付-发货-收货流程验证、优惠券原子领券验证、购物车商品归属验证、地址归属校验验证、评价订单关联验证、收藏幂等验证、秒杀活动状态机验证、数据库事务一致性验证。')
doc.add_page_break()
H('9.5 Postman 集合设置', level=2)
P('在Postman中依次导入以下集合和变量即可运行：')
B('步骤一：导入Postman集合')
P('打开Postman → File → Import → 选择 docs/postman_collection.json 导入。')
B('步骤二：设置环境变量')
P('点击"Environments" → 添加新环境（如"FlashSale Local"），设置以下变量：')
T(['变量','值','说明'],[['base_url','http://localhost:8080','Gateway地址'],['admin_token','<管理员JWT>','管理员登录后获取'],['customer_token','<用户JWT>','用户登录后获取'],['test_activity_id','1','秒杀活动ID']])
B('步骤三：运行顺序')
P('推荐按以下顺序手动执行（也可使用Postman Runner自动执行）：')
[doc.add_paragraph(i, style='List Bullet') for i in [
    '1. 用户注册/登录（获取token）',
    '2. 商品浏览、搜索',
    '3. 地址管理',
    '4. 购物车操作',
    '5. 下单与支付',
    '6. 秒杀活动管理',
    '7. 秒杀抢购',
    '8. 优惠券领取与使用',
    '9. 管理后台操作',
    '10. 数据报表查看']]
H('9.6 万人并发演示数据', level=2)
P('万人并发演示脚本位于 e2e/ 目录：')
[doc.add_paragraph(i, style='List Bullet') for i in [
    'wanren_tongchang.py — 50品万人同场·10000单（御准方案），三模式切换 --mode seckill|shopping|mixed，覆盖秒杀/普通购物/混跑绝对E2E流程',
    'stress_validation_v2.py — 架构验证：11项正确性断言']]
P('运行大型演示前建议先执行 stress_validation_v2.py 验证基础架构正确性，确保库存扣减和订单状态机均无误。')

# ============ CH10 (was CH9) ============
doc.add_page_break()
H('十、生产环境演进方案', level=1)
H('10.1 高可用（HA）', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in [
    'MySQL主从+读写分离（ProxySQL或MyCat中间件）',
    'Redis Sentinel 3节点哨兵，自动切换',
    'RabbitMQ镜像队列（Mirror Queue），queue同步到所有节点',
    'Gateway + Nginx双活部署，Keepalived浮动IP']]
H('10.2 扩容', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in [
    '水平扩容：所有服务无状态（JWT），加实例即可',
    '分库分表：ShardingSphere按user_id分库，order_id分表',
    'CDN静态资源：前端资源部署CDN，Gateway只处理API',
    'Redis Cluster：3主3从分片集群，支撑更大库存Key并发']]
H('10.3 LLM优化', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in [
    '异步非阻塞：WebClient调用LLM + CompletableFuture回调',
    '请求排队：Redis队列控制并发不超过API限流（60 RPM）',
    '相似缓存：缓存重复问题回复（语义相似度>0.95）',
    '流式推送：改为SSE逐步返回']]

# ============ CH11 (was CH10) ============
doc.add_page_break()
H('十一、AI辅助设计开发总结', level=1)
H('11.1 开发过程', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in [
    '需求分析与系统设计阶段：AI协助生成架构图、ER图、接口设计',
    '编码实现阶段：AI生成代码框架、Mapper接口、Controller模板',
    '测试阶段：AI自动生成E2E测试脚本、并发测试脚本、性能分析',
    '文档编写阶段：AI生成项目文档、操作手册、报告']]
H('11.2 AI与人类的协作模式', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in [
    'AI负责：代码生成、测试脚本编写、文档生成、重复性任务',
    '人类负责：架构决策、技术选型、关键代码审查、系统设计权衡',
    '协作亮点：测试脚本全程由AI编写并验证，人类只审查边界条件',
    '工具链：DeepSeek(Chat) + Claude(arch) + QwenPaw 三Agent协作']]
H('11.3 关键决策记录', level=2)
T(['#','决策','AI建议','最终选择','理由'],[['1','WebSocket路由归属','seckill-service','order-service','sessionMap在order-service，推送路径短'],['2','Redis vs Lua','Lua保证原子性','DECR即可','单key操作无需脚本，性能更高'],['3','DB直插 vs HTTP','DB直插性能高','HTTP全流程','E2E必须走HTTP验证全链路']])
H('11.4 经验与反思', level=2)
[doc.add_paragraph(i, style='List Bullet') for i in [
    '数据库直插vs HTTP注册：性能差几十倍，但E2E测试必须走HTTP才能验证全链路',
    'Gateway限流配置：令牌桶容量和填充速率需要根据实际压测数据调整',
    'WebSocket实时推送在多服务架构下需要统一规划连接归属',
    'AI辅助开发大幅提升了编码效率，但架构设计和技术决策仍需人类判断']]

# ============ MERGE & SAVE ============
out = '/tmp/report_merged.docx'
doc.save(out)
print(f'Merged report saved to {out}')
